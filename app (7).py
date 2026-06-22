"""
ระบบจัดการไซต์งาน — Streamlit + Google Sheets + Google Drive (Service Account)
"""

import streamlit as st
import pandas as pd
import json
import os
import io
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches
from PIL import Image as PILImage

import base64
import requests
import gspread
from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload, MediaIoBaseDownload

# ─────────────────────────────────────────────────────────────────────────────
# 0. ตั้งค่า
# ─────────────────────────────────────────────────────────────────────────────
KEY_FILE = os.environ.get("GOOGLE_KEY_FILE", "key.json")

SCOPES = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive",
]
SHEET_NAME = "Sites"
SHEET_HEADER = ["SiteName", "Status", "CreatedAt", "StartMile", "EndMile", "DataJSON"]

HOTEL_COUNT = 3
HOTEL_ITEM_COUNT = 3
FUEL_COUNT = 20

st.set_page_config(page_title="ระบบจัดการไซต์งาน", layout="wide")


# ─────────────────────────────────────────────────────────────────────────────
# 1. Google API clients — โหลด credentials จาก Secrets หรือ key.json
# ─────────────────────────────────────────────────────────────────────────────
def _get_secret(key: str, env_key: str = "") -> str:
    try:
        v = st.secrets.get(key, "")
        return v if v else ""
    except Exception:
        return os.environ.get(env_key or key, "")


SPREADSHEET_ID  = _get_secret("SPREADSHEET_ID",  "GOOGLE_SHEET_ID")
DRIVE_FOLDER_ID = _get_secret("DRIVE_FOLDER_ID", "GOOGLE_DRIVE_FOLDER_ID")


@st.cache_resource
def _get_credentials():
    """โหลด credentials:
    1. Streamlit Secrets [gcp_service_account]  (Streamlit Cloud)
    2. ไฟล์ key.json  (รันในเครื่อง)
    """
    # ── 1. Streamlit Secrets ──────────────────────────────────────────────
    try:
        sa_info = dict(st.secrets["gcp_service_account"])
        return Credentials.from_service_account_info(sa_info, scopes=SCOPES)
    except (KeyError, FileNotFoundError):
        pass
    except Exception as e:
        raise RuntimeError(f"โหลด credentials จาก Streamlit Secrets ไม่สำเร็จ: {e}")

    # ── 2. ไฟล์ key.json ─────────────────────────────────────────────────
    if not os.path.exists(KEY_FILE):
        raise RuntimeError(
            "ไม่พบ credentials\n"
            "• Streamlit Cloud: ไปที่ App settings → Secrets ใส่ [gcp_service_account]\n"
            f"• Local: วางไฟล์ key.json ไว้ในโฟลเดอร์เดียวกับ app.py"
        )
    return Credentials.from_service_account_file(KEY_FILE, scopes=SCOPES)


@st.cache_resource
def _get_gspread_client():
    return gspread.authorize(_get_credentials())


@st.cache_resource
def _get_drive_service():
    return build("drive", "v3", credentials=_get_credentials())


def _get_worksheet():
    """เปิด worksheet 'Sites' — สร้างให้อัตโนมัติถ้ายังไม่มี"""
    if not SPREADSHEET_ID:
        st.error("ยังไม่ได้ตั้งค่า SPREADSHEET_ID")
        st.stop()
    gc = _get_gspread_client()
    sh = gc.open_by_key(SPREADSHEET_ID)
    try:
        ws = sh.worksheet(SHEET_NAME)
    except gspread.WorksheetNotFound:
        ws = sh.add_worksheet(title=SHEET_NAME, rows=200, cols=len(SHEET_HEADER))
        ws.append_row(SHEET_HEADER)
    if not ws.row_values(1):
        ws.update([SHEET_HEADER], value_input_option="RAW")
    return ws


# ─────────────────────────────────────────────────────────────────────────────
# 2. การจัดการข้อมูลไซต์
# ─────────────────────────────────────────────────────────────────────────────
def _empty_site_data() -> dict:
    return {
        "hotels": {
            str(h): {str(i): {"img": "", "desc": ""} for i in range(1, HOTEL_ITEM_COUNT + 1)}
            for h in range(1, HOTEL_COUNT + 1)
        },
        "car_img": "",
        "mile_start_img": "",
        "mile_end_img": "",
        "fuel": {
            str(n): {"bill": "", "pre": "", "post": "", "date": "", "province": ""}
            for n in range(1, FUEL_COUNT + 1)
        },
    }


def load_history() -> pd.DataFrame:
    try:
        ws = _get_worksheet()
        records = ws.get_all_records()
        if not records:
            return pd.DataFrame(columns=SHEET_HEADER)
        df = pd.DataFrame(records)
        for c in SHEET_HEADER:
            if c not in df.columns:
                df[c] = ""
        return df[SHEET_HEADER].astype(str)
    except Exception as e:
        st.warning(f"โหลดข้อมูลจาก Google Sheet ไม่สำเร็จ: {e}")
        return pd.DataFrame(columns=SHEET_HEADER)


def _save_history(df: pd.DataFrame):
    ws = _get_worksheet()
    rows = [SHEET_HEADER] + df[SHEET_HEADER].astype(str).values.tolist()
    ws.clear()
    ws.update(rows, value_input_option="RAW")


def _load_site_data(site_name: str) -> dict:
    df = load_history()
    row = df[df["SiteName"] == site_name]
    if row.empty:
        return _empty_site_data()
    raw = row.iloc[0].get("DataJSON", "")
    if not raw:
        return _empty_site_data()
    try:
        data = json.loads(raw)
        default = _empty_site_data()
        for k, v in default.items():
            data.setdefault(k, v)
        return data
    except Exception:
        return _empty_site_data()


def _save_site_data(site_name: str, site_data: dict):
    df = load_history()
    if site_name not in df["SiteName"].values:
        new_row = {
            "SiteName": site_name, "Status": "Draft",
            "CreatedAt": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M"),
            "StartMile": "0", "EndMile": "0",
            "DataJSON": json.dumps(site_data, ensure_ascii=False),
        }
        df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)
    else:
        idx = df[df["SiteName"] == site_name].index[0]
        df.at[idx, "DataJSON"] = json.dumps(site_data, ensure_ascii=False)
    _save_history(df)


def create_or_open_site(site_name: str) -> str:
    site_name = (site_name or "").strip()
    if not site_name:
        st.error("กรุณาระบุชื่อไซต์งาน")
        return ""
    df = load_history()
    if site_name not in df["SiteName"].values:
        new_row = {
            "SiteName": site_name, "Status": "Draft",
            "CreatedAt": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M"),
            "StartMile": "0", "EndMile": "0",
            "DataJSON": json.dumps(_empty_site_data(), ensure_ascii=False),
        }
        df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)
        _save_history(df)
    return site_name


# ─────────────────────────────────────────────────────────────────────────────
# 3. อัปโหลดรูปขึ้น Google Drive
# ─────────────────────────────────────────────────────────────────────────────
_drive_folder_cache: dict = {}


def _get_or_create_site_folder(site_name: str) -> str:
    if site_name in _drive_folder_cache:
        return _drive_folder_cache[site_name]
    if not DRIVE_FOLDER_ID:
        st.error("ยังไม่ได้ตั้งค่า DRIVE_FOLDER_ID")
        st.stop()
    service = _get_drive_service()
    safe_name = "".join(c for c in site_name if c not in '\\/:*?"<>|').strip()
    query = (
        f"name = '{safe_name}' and mimeType = 'application/vnd.google-apps.folder' "
        f"and '{DRIVE_FOLDER_ID}' in parents and trashed = false"
    )
    results = service.files().list(q=query, fields="files(id, name)", supportsAllDrives=True, includeItemsFromAllDrives=True, corpora="allDrives").execute()
    existing = results.get("files", [])
    if existing:
        folder_id = existing[0]["id"]
    else:
        meta = {
            "name": safe_name,
            "mimeType": "application/vnd.google-apps.folder",
            "parents": [DRIVE_FOLDER_ID],
        }
        created = service.files().create(body=meta, fields="id", supportsAllDrives=True).execute()
        folder_id = created["id"]
    _drive_folder_cache[site_name] = folder_id
    return folder_id


def _upload_image(site_name: str, uploaded_file, dest_filename: str) -> str:
    """อัปโหลดรูปไป Cloudinary (ฟรี 25GB ไม่ต้อง OAuth)"""
    if not uploaded_file:
        return ""
    try:
        cloud_name = _get_secret("CLOUDINARY_CLOUD_NAME", "CLOUDINARY_CLOUD_NAME")
        api_key    = _get_secret("CLOUDINARY_API_KEY",    "CLOUDINARY_API_KEY")
        api_secret = _get_secret("CLOUDINARY_API_SECRET", "CLOUDINARY_API_SECRET")
        if not all([cloud_name, api_key, api_secret]):
            st.warning("ยังไม่ได้ตั้งค่า Cloudinary ใน Streamlit Secrets")
            return ""

        with PILImage.open(uploaded_file) as im:
            if im.mode in ("RGBA", "P", "LA"):
                im = im.convert("RGB")
            buf = io.BytesIO()
            im.save(buf, format="JPEG", quality=85)
            buf.seek(0)

        import hashlib, time
        timestamp = str(int(time.time()))
        public_id = f"trip/{site_name}/{dest_filename}"
        # สร้าง signature
        sig_str = f"public_id={public_id}&timestamp={timestamp}{api_secret}"
        signature = hashlib.sha1(sig_str.encode()).hexdigest()

        resp = requests.post(
            f"https://api.cloudinary.com/v1_1/{cloud_name}/image/upload",
            data={
                "api_key":   api_key,
                "timestamp": timestamp,
                "public_id": public_id,
                "signature": signature,
            },
            files={"file": ("image.jpg", buf, "image/jpeg")},
            timeout=60,
        )
        resp.raise_for_status()
        return resp.json().get("secure_url", "")

    except Exception as e:
        st.warning(f"อัปโหลดรูปไม่สำเร็จ ({dest_filename}): {e}")
        return ""

    try:
        with PILImage.open(uploaded_file) as im:
            if im.mode in ("RGBA", "P", "LA"):
                im = im.convert("RGB")
            buf = io.BytesIO()
            im.save(buf, format="JPEG", quality=85)
            img_bytes = buf.getvalue()

        token = _get_cached_token()
        headers = {
            "Authorization": f"Bearer {token}",
            "Content-Type": "image/jpeg",
        }

        onedrive_folder = _get_secret("ONEDRIVE_FOLDER", "ONEDRIVE_FOLDER") or "TripPhotos"
        filename = f"{dest_filename}.jpg"
        # อัปโหลดตรงไปที่ path /TripPhotos/{site_name}/{filename}
        safe_site = "".join(c for c in site_name if c not in r'\/:*?"<>|').strip()
        upload_url = (
            f"https://graph.microsoft.com/v1.0/me/drive/root:/"
            f"{onedrive_folder}/{safe_site}/{filename}:/content"
        )
        resp = requests.put(upload_url, headers=headers, data=img_bytes, timeout=60)
        resp.raise_for_status()
        data = resp.json()

        # ทำให้ลิงก์เปิดได้สาธารณะ
        file_id = data["id"]
        share_url = f"https://graph.microsoft.com/v1.0/me/drive/items/{file_id}/createLink"
        share_resp = requests.post(share_url,
            headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"},
            json={"type": "view", "scope": "anonymous"},
            timeout=15,
        )
        if share_resp.status_code in (200, 201):
            link = share_resp.json().get("link", {}).get("webUrl", "")
            # แปลง share link เป็น direct image URL
            if "1drv.ms" in link or "sharepoint.com" in link:
                return link
        # fallback: ใช้ webUrl ตรงๆ
        return data.get("webUrl", "")

    except Exception as e:
        st.warning(f"อัปโหลดรูปไม่สำเร็จ ({dest_filename}): {e}")
        return ""


# ─────────────────────────────────────────────────────────────────────────────
# 4. สร้าง Word / ดาวน์โหลดรูปจาก Drive
# ─────────────────────────────────────────────────────────────────────────────
def _fetch_drive_image_bytes(url: str):
    if not url:
        return None
    try:
        m = re.search(r'[?&]id=([^&]+)', url)
        if not m:
            return None
        file_id = m.group(1)
        service = _get_drive_service()
        request = service.files().get_media(fileId=file_id)
        buf = io.BytesIO()
        downloader = MediaIoBaseDownload(buf, request)
        done = False
        while not done:
            _, done = downloader.next_chunk()
        buf.seek(0)
        return buf.read()
    except Exception as e:
        print(f"[เตือน] ดาวน์โหลดรูปจาก Drive ไม่สำเร็จ ({url}): {e}")
        return None


def _safe_add_picture(target, img_url, width):
    if not img_url:
        return False
    raw = _fetch_drive_image_bytes(img_url)
    if not raw:
        return False
    try:
        with PILImage.open(io.BytesIO(raw)) as im:
            if im.mode in ("RGBA", "P", "LA"):
                im = im.convert("RGB")
            buf = io.BytesIO()
            im.save(buf, format="JPEG", quality=90)
            buf.seek(0)

        if hasattr(target, "add_picture"):
            target.add_picture(buf, width=width)
        else:
            p = target.add_paragraph()
            run = p.add_run()
            run.add_picture(buf, width=width)
        return True
    except Exception as e:
        print(f"[เตือน] แทรกรูปไม่สำเร็จ ({img_url}): {e}")
        return False


def build_word_report(site_name: str, site_data: dict, mile_start: float, mile_end: float) -> io.BytesIO:
    doc = Document()
    doc.add_heading(f'Trip Report — {site_name}', 0)
    doc.add_paragraph(f"สร้างเมื่อ: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    doc.add_heading('ข้อมูลรถและการเดินทาง', level=1)
    distance = max(0, mile_end - mile_start)
    doc.add_paragraph(f"เลขไมล์เริ่มต้น: {mile_start:.0f}")
    doc.add_paragraph(f"เลขไมล์หลังจบ: {mile_end:.0f}")
    doc.add_paragraph(f"ระยะทางรวม: {distance:.0f} กม.")

    car_img = site_data.get("car_img", "")
    mile_start_img = site_data.get("mile_start_img", "")
    mile_end_img = site_data.get("mile_end_img", "")
    if car_img or mile_start_img or mile_end_img:
        car_table = doc.add_table(rows=1, cols=2)
        car_table.autofit = False
        car_table.columns[0].width = Inches(3.5)
        car_table.columns[1].width = Inches(3.0)
        left_cell = car_table.cell(0, 0)
        left_cell.paragraphs[0].text = "รูปรถ"
        if car_img:
            if not _safe_add_picture(left_cell, car_img, Inches(3.2)):
                left_cell.add_paragraph("(ไม่สามารถแสดงรูปนี้ได้)")
        else:
            left_cell.add_paragraph("(ไม่มีรูปรถ)")
        right_cell = car_table.cell(0, 1)
        right_cell.paragraphs[0].text = "ไมล์เริ่มต้น"
        if mile_start_img:
            if not _safe_add_picture(right_cell, mile_start_img, Inches(2.2)):
                right_cell.add_paragraph("(ไม่สามารถแสดงรูปนี้ได้)")
        else:
            right_cell.add_paragraph("(ไม่มีรูป)")
        right_cell.add_paragraph("ไมล์หลังจบ")
        if mile_end_img:
            if not _safe_add_picture(right_cell, mile_end_img, Inches(2.2)):
                right_cell.add_paragraph("(ไม่สามารถแสดงรูปนี้ได้)")
        else:
            right_cell.add_paragraph("(ไม่มีรูป)")
        doc.add_paragraph("")

    doc.add_heading('รูปภาพโรงแรม', level=1)
    any_hotel = False
    for h in range(1, HOTEL_COUNT + 1):
        items = site_data["hotels"][str(h)]
        has_any_img = any(items[str(i)]["img"] for i in range(1, HOTEL_ITEM_COUNT + 1))
        if not has_any_img:
            continue
        any_hotel = True
        doc.add_heading(f'โรงแรมที่ {h}', level=2)
        hotel_table = doc.add_table(rows=2, cols=HOTEL_ITEM_COUNT)
        hotel_table.autofit = True
        img_row = hotel_table.rows[0]
        desc_row = hotel_table.rows[1]
        for i in range(1, HOTEL_ITEM_COUNT + 1):
            item = items[str(i)]
            img_cell = img_row.cells[i - 1]
            desc_cell = desc_row.cells[i - 1]
            if item["img"]:
                if not _safe_add_picture(img_cell, item["img"], Inches(1.9)):
                    img_cell.paragraphs[0].text = "(ไม่สามารถแสดงรูปนี้ได้)"
            else:
                img_cell.paragraphs[0].text = "(ไม่มีรูป)"
            desc_cell.paragraphs[0].text = item["desc"] or "(ไม่มีคำอธิบาย)"
        doc.add_paragraph("")
    if not any_hotel:
        doc.add_paragraph("(ยังไม่มีรูปโรงแรมที่บันทึกแล้ว)")

    doc.add_heading('บันทึกการเติมน้ำมัน', level=1)
    any_fuel = False
    for n in range(1, FUEL_COUNT + 1):
        f = site_data["fuel"][str(n)]
        if not (f["bill"] or f["pre"] or f["post"]):
            continue
        any_fuel = True
        doc.add_heading(f'การเติมครั้งที่ {n}', level=2)
        doc.add_paragraph(f"วันที่: {f['date'] or '-'}   จังหวัด: {f['province'] or '-'}")
        fuel_table = doc.add_table(rows=1, cols=2)
        fuel_table.autofit = False
        fuel_table.columns[0].width = Inches(4.0)
        fuel_table.columns[1].width = Inches(2.5)
        left_cell = fuel_table.cell(0, 0)
        left_cell.paragraphs[0].text = "ใบเสร็จ"
        if f["bill"]:
            if not _safe_add_picture(left_cell, f["bill"], Inches(3.7)):
                left_cell.add_paragraph("(ไม่สามารถแสดงรูปนี้ได้)")
        else:
            left_cell.add_paragraph("(ไม่มีรูปใบเสร็จ)")
        right_cell = fuel_table.cell(0, 1)
        right_cell.paragraphs[0].text = "ไมล์ก่อนเติม"
        if f["pre"]:
            if not _safe_add_picture(right_cell, f["pre"], Inches(2.2)):
                right_cell.add_paragraph("(ไม่สามารถแสดงรูปนี้ได้)")
        else:
            right_cell.add_paragraph("(ไม่มีรูป)")
        right_cell.add_paragraph("ไมล์หลังเติม")
        if f["post"]:
            if not _safe_add_picture(right_cell, f["post"], Inches(2.2)):
                right_cell.add_paragraph("(ไม่สามารถแสดงรูปนี้ได้)")
        else:
            right_cell.add_paragraph("(ไม่มีรูป)")
        doc.add_paragraph("")
    if not any_fuel:
        doc.add_paragraph("(ยังไม่มีรายการเติมน้ำมันที่บันทึกแล้ว)")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def upload_report_to_drive(site_name: str, doc_buf: io.BytesIO) -> str:
    try:
        service = _get_drive_service()
        folder_id = _get_or_create_site_folder(site_name)
        media = MediaIoBaseUpload(
            doc_buf,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            resumable=False,
        )
        meta = {"name": f"{site_name}_Report.docx", "parents": [folder_id]}
        query = (
            f"name = '{site_name}_Report.docx' and '{folder_id}' in parents "
            f"and trashed = false"
        )
        existing = service.files().list(q=query, fields="files(id)", supportsAllDrives=True, includeItemsFromAllDrives=True, corpora="allDrives").execute().get("files", [])
        if existing:
            file_id = existing[0]["id"]
            service.files().update(fileId=file_id, media_body=media, supportsAllDrives=True).execute()
        else:
            created = service.files().create(body=meta, media_body=media, fields="id", supportsAllDrives=True).execute()
            file_id = created["id"]
            service.permissions().create(
                fileId=file_id, body={"type": "anyone", "role": "reader"}
            ).execute()
        return f"https://drive.google.com/file/d/{file_id}/view"
    except Exception as e:
        st.warning(f"อัปโหลดรายงานไม่สำเร็จ: {e}")
        return ""


# ─────────────────────────────────────────────────────────────────────────────
# 5. Session state helpers
# ─────────────────────────────────────────────────────────────────────────────
def _init_session():
    if "page" not in st.session_state:
        st.session_state.page = "home"
    if "current_site" not in st.session_state:
        st.session_state.current_site = ""
    if "site_data" not in st.session_state:
        st.session_state.site_data = _empty_site_data()


# ─────────────────────────────────────────────────────────────────────────────
# 6. หน้าแรก — รายการไซต์งาน
# ─────────────────────────────────────────────────────────────────────────────
def page_home():
    st.title("🏗️ ระบบจัดการไซต์งาน")

    st.subheader("ประวัติไซต์งาน")
    df = load_history()
    if df.empty:
        st.info("ยังไม่มีไซต์งาน กรุณาสร้างไซต์งานใหม่ด้านล่าง")
    else:
        display_df = df[["SiteName", "Status", "CreatedAt"]].copy()
        display_df.columns = ["ชื่อไซต์", "สถานะ", "วันที่สร้าง"]
        st.dataframe(display_df, use_container_width=True, hide_index=True)

        st.write("**เลือกไซต์ที่ต้องการแก้ไข:**")
        site_names = df["SiteName"].tolist()
        selected = st.selectbox("ไซต์งาน", site_names, label_visibility="collapsed")
        if st.button("✏️ แก้ไขไซต์ที่เลือก", type="secondary"):
            st.session_state.current_site = selected
            st.session_state.site_data = _load_site_data(selected)
            st.session_state.page = "entry"
            st.rerun()

    st.divider()
    st.subheader("สร้างไซต์งานใหม่")
    col1, col2 = st.columns([3, 1])
    with col1:
        new_name = st.text_input("ชื่อไซต์งานใหม่", label_visibility="collapsed",
                                  placeholder="ระบุชื่อไซต์งาน...")
    with col2:
        if st.button("➕ สร้าง", type="primary", use_container_width=True):
            if not new_name.strip():
                st.error("กรุณาระบุชื่อไซต์งาน")
            else:
                name = create_or_open_site(new_name.strip())
                if name:
                    st.session_state.current_site = name
                    st.session_state.site_data = _load_site_data(name)
                    st.session_state.page = "entry"
                    st.rerun()


# ─────────────────────────────────────────────────────────────────────────────
# 7. หน้าบันทึกข้อมูล
# ─────────────────────────────────────────────────────────────────────────────
def page_entry():
    site_name = st.session_state.current_site
    site_data = st.session_state.site_data

    col_back, col_title = st.columns([1, 5])
    with col_back:
        if st.button("⬅️ กลับหน้าแรก"):
            st.session_state.page = "home"
            st.rerun()
    with col_title:
        st.title(f"📋 {site_name}")

    # ─── 1. ข้อมูลโรงแรม ───
    with st.expander("1. ข้อมูลโรงแรม", expanded=False):
        for h in range(1, HOTEL_COUNT + 1):
            st.markdown(f"#### โรงแรมที่ {h}")
            cols = st.columns(HOTEL_ITEM_COUNT)
            new_imgs = {}
            for i in range(1, HOTEL_ITEM_COUNT + 1):
                with cols[i - 1]:
                    cur_img_url = site_data["hotels"][str(h)][str(i)]["img"]
                    if cur_img_url:
                        st.image(cur_img_url, caption=f"รูปปัจจุบัน {i}", use_container_width=True)
                    uploaded = st.file_uploader(
                        f"อัปโหลดรูป {i}",
                        type=["jpg", "jpeg", "png", "webp", "heic"],
                        key=f"hotel_{h}_{i}_img",
                    )
                    new_imgs[i] = uploaded
                    site_data["hotels"][str(h)][str(i)]["desc"] = st.text_area(
                        f"คำอธิบายรูป {i}",
                        value=site_data["hotels"][str(h)][str(i)]["desc"],
                        key=f"hotel_{h}_{i}_desc",
                        height=80,
                    )

            if st.button(f"💾 บันทึกโรงแรมที่ {h}", key=f"save_hotel_{h}"):
                with st.spinner("กำลังอัปโหลดและบันทึก..."):
                    for i in range(1, HOTEL_ITEM_COUNT + 1):
                        if new_imgs[i]:
                            url = _upload_image(site_name, new_imgs[i], f"hotel_{h}_{i}")
                            if url:
                                site_data["hotels"][str(h)][str(i)]["img"] = url
                    _save_site_data(site_name, site_data)
                    st.session_state.site_data = site_data
                st.success(f"✅ บันทึกโรงแรมที่ {h} เรียบร้อย")

            st.divider()

    # ─── 2. ข้อมูลรถ ───
    with st.expander("2. ข้อมูลรถ", expanded=False):
        df = load_history()
        row_data = df[df["SiteName"] == site_name].iloc[0] if site_name in df["SiteName"].values else {}
        mile_start_val = float(row_data.get("StartMile", 0) or 0)
        mile_end_val = float(row_data.get("EndMile", 0) or 0)

        col_car, col_mile = st.columns([2, 1])
        with col_car:
            if site_data.get("car_img"):
                st.image(site_data["car_img"], caption="รูปรถปัจจุบัน", use_container_width=True)
            car_file = st.file_uploader("อัปโหลดรูปรถ", type=["jpg", "jpeg", "png", "webp"], key="car_img")
        with col_mile:
            if site_data.get("mile_start_img"):
                st.image(site_data["mile_start_img"], caption="รูปไมล์เริ่ม", use_container_width=True)
            mile_start_file = st.file_uploader("รูปไมล์เริ่ม", type=["jpg", "jpeg", "png", "webp"], key="mile_start_img")
            if site_data.get("mile_end_img"):
                st.image(site_data["mile_end_img"], caption="รูปไมล์จบ", use_container_width=True)
            mile_end_file = st.file_uploader("รูปไมล์จบ", type=["jpg", "jpeg", "png", "webp"], key="mile_end_img")

        col_m1, col_m2 = st.columns(2)
        with col_m1:
            mile_start = st.number_input("เลขไมล์เริ่มต้น", value=mile_start_val, key="mile_start")
        with col_m2:
            mile_end = st.number_input("เลขไมล์หลังจบ", value=mile_end_val, key="mile_end")

        if st.button("💾 บันทึกข้อมูลรถ", key="save_car"):
            with st.spinner("กำลังอัปโหลดและบันทึก..."):
                if car_file:
                    url = _upload_image(site_name, car_file, "car")
                    if url:
                        site_data["car_img"] = url
                if mile_start_file:
                    url = _upload_image(site_name, mile_start_file, "mile_start")
                    if url:
                        site_data["mile_start_img"] = url
                if mile_end_file:
                    url = _upload_image(site_name, mile_end_file, "mile_end")
                    if url:
                        site_data["mile_end_img"] = url

                df2 = load_history()
                idx = df2[df2["SiteName"] == site_name].index
                if len(idx) > 0:
                    df2.at[idx[0], "StartMile"] = str(mile_start or 0)
                    df2.at[idx[0], "EndMile"] = str(mile_end or 0)
                    _save_history(df2)
                _save_site_data(site_name, site_data)
                st.session_state.site_data = site_data
            st.success("✅ บันทึกข้อมูลรถเรียบร้อย")

    # ─── 3. การเติมน้ำมัน ───
    with st.expander(f"3. การเติมน้ำมัน ({FUEL_COUNT} ครั้ง)", expanded=False):
        for n in range(1, FUEL_COUNT + 1):
            with st.expander(f"เติมครั้งที่ {n}", expanded=False):
                f = site_data["fuel"][str(n)]
                col_bill, col_mile2 = st.columns([2, 1])
                with col_bill:
                    if f.get("bill"):
                        st.image(f["bill"], caption="ใบเสร็จปัจจุบัน", use_container_width=True)
                    bill_file = st.file_uploader("ใบเสร็จ", type=["jpg", "jpeg", "png", "webp"],
                                                  key=f"fuel_{n}_bill")
                with col_mile2:
                    if f.get("pre"):
                        st.image(f["pre"], caption="ไมล์ก่อนเติม", use_container_width=True)
                    pre_file = st.file_uploader("ไมล์ก่อนเติม", type=["jpg", "jpeg", "png", "webp"],
                                                 key=f"fuel_{n}_pre")
                    if f.get("post"):
                        st.image(f["post"], caption="ไมล์หลังเติม", use_container_width=True)
                    post_file = st.file_uploader("ไมล์หลังเติม", type=["jpg", "jpeg", "png", "webp"],
                                                  key=f"fuel_{n}_post")

                col_d, col_p = st.columns(2)
                with col_d:
                    date_val = st.text_input("วันที่เติม (เช่น 2026-06-21)",
                                              value=f.get("date", ""), key=f"fuel_{n}_date")
                with col_p:
                    prov_val = st.text_input("จังหวัด",
                                              value=f.get("province", ""), key=f"fuel_{n}_prov")

                if st.button(f"💾 บันทึกการเติมครั้งที่ {n}", key=f"save_fuel_{n}"):
                    with st.spinner("กำลังอัปโหลดและบันทึก..."):
                        if bill_file:
                            url = _upload_image(site_name, bill_file, f"fuel_{n}_bill")
                            if url:
                                site_data["fuel"][str(n)]["bill"] = url
                        if pre_file:
                            url = _upload_image(site_name, pre_file, f"fuel_{n}_pre")
                            if url:
                                site_data["fuel"][str(n)]["pre"] = url
                        if post_file:
                            url = _upload_image(site_name, post_file, f"fuel_{n}_post")
                            if url:
                                site_data["fuel"][str(n)]["post"] = url
                        site_data["fuel"][str(n)]["date"] = date_val or ""
                        site_data["fuel"][str(n)]["province"] = prov_val or ""
                        _save_site_data(site_name, site_data)
                        st.session_state.site_data = site_data
                    st.success(f"✅ บันทึกการเติมน้ำมันครั้งที่ {n} เรียบร้อย")

    # ─── Export ───
    st.divider()
    st.subheader("ยืนยันและสร้างรายงาน")
    col_exp1, col_exp2 = st.columns([1, 1])
    with col_exp1:
        if st.button("✅ ยืนยันและอัปโหลดรายงาน Word ลง Drive", type="primary", use_container_width=True):
            df3 = load_history()
            row3 = df3[df3["SiteName"] == site_name]
            ms = float(row3.iloc[0].get("StartMile", 0) or 0) if not row3.empty else 0
            me = float(row3.iloc[0].get("EndMile", 0) or 0) if not row3.empty else 0

            with st.spinner("กำลังสร้างและอัปโหลดรายงาน..."):
                doc_buf = build_word_report(site_name, site_data, ms, me)
                df3.loc[df3["SiteName"] == site_name, "Status"] = "Confirmed"
                _save_history(df3)
                doc_buf_for_drive = io.BytesIO(doc_buf.getvalue())
                report_url = upload_report_to_drive(site_name, doc_buf_for_drive)
                st.session_state["last_report_buf"] = doc_buf.getvalue()
                st.session_state["last_report_url"] = report_url

            if report_url:
                st.success(f"✅ ยืนยันสำเร็จ — [เปิดรายงานใน Drive]({report_url})")
            else:
                st.warning("⚠️ ยืนยันสำเร็จ แต่อัปโหลดรายงานขึ้น Drive ไม่สำเร็จ")

    with col_exp2:
        if "last_report_buf" in st.session_state:
            st.download_button(
                label="⬇️ ดาวน์โหลดไฟล์ Word",
                data=st.session_state["last_report_buf"],
                file_name=f"{site_name}_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        else:
            df4 = load_history()
            row4 = df4[df4["SiteName"] == site_name]
            ms2 = float(row4.iloc[0].get("StartMile", 0) or 0) if not row4.empty else 0
            me2 = float(row4.iloc[0].get("EndMile", 0) or 0) if not row4.empty else 0
            doc_preview = build_word_report(site_name, site_data, ms2, me2)
            st.download_button(
                label="⬇️ ดาวน์โหลดไฟล์ Word (ร่าง)",
                data=doc_preview.getvalue(),
                file_name=f"{site_name}_Report_draft.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )


# ─────────────────────────────────────────────────────────────────────────────
# 8. Main
# ─────────────────────────────────────────────────────────────────────────────
def main():
    _init_session()
    if st.session_state.page == "home":
        page_home()
    elif st.session_state.page == "entry":
        page_entry()


if __name__ == "__main__":
    main()
