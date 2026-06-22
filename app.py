"""
ระบบจัดการไซต์งาน — Gradio + Google Sheets + Google Drive (Service Account)
═══════════════════════════════════════════════════════════════════════
ออกแบบมาให้รันบนเว็บ (เช่น Streamlit Cloud, Render, Hugging Face Spaces ฯลฯ)
ไม่ใช่ Google Colab — ใช้ Service Account (ไฟล์ key.json) แทนการ
drive.mount() แบบ interactive เพราะรันบนเว็บไม่มีคนนั่งกด popup Allow

ข้อมูลทั้งหมดเก็บถาวรภายนอกตัวแอป (ไม่ใช่ local disk ของเซิร์ฟเวอร์ที่
อาจหายเวลา redeploy/restart):
  - ข้อมูล metadata (สถานะ/วันที่/ไมล์/คำอธิบาย ฯลฯ) → Google Sheet
  - รูปภาพทั้งหมด (โรงแรม/รถ/น้ำมัน) → Google Drive

═══════════════════════════════════════════════════════════════════════
การตั้งค่าที่ต้องทำก่อนใช้งาน (ทำครั้งเดียว)
═══════════════════════════════════════════════════════════════════════
1. สร้าง Service Account ใน Google Cloud Console:
   - ไปที่ console.cloud.google.com → เลือก/สร้างโปรเจกต์
   - เปิดใช้งาน "Google Sheets API" และ "Google Drive API"
   - ไปที่ APIs & Services → Credentials → Create Credentials
     → Service Account → สร้างเสร็จแล้วเข้าไปที่ Service Account นั้น
     → แท็บ Keys → Add Key → Create new key → เลือก JSON
   - ไฟล์ JSON จะถูกดาวน์โหลดมา เปลี่ยนชื่อเป็น "key.json"
     แล้ววางไว้ในโฟลเดอร์เดียวกับไฟล์ app.py นี้
     (หรือกำหนด path อื่นผ่าน environment variable GOOGLE_KEY_FILE)

2. แชร์สิทธิ์ให้ Service Account เข้าถึงได้:
   - เปิดไฟล์ key.json ดูค่า "client_email" (รูปแบบ xxx@xxx.iam.gserviceaccount.com)
   - สร้าง Google Sheet ใหม่ 1 ไฟล์ (ชื่ออะไรก็ได้) → กด "แชร์" →
     วาง client_email ลงไป → ตั้งสิทธิ์เป็น "ผู้แก้ไข (Editor)"
   - คัดลอก Spreadsheet ID จาก URL ของ Sheet นั้น
     (ส่วนระหว่าง /d/ กับ /edit) ใส่ในตัวแปร SPREADSHEET_ID ด้านล่าง
   - สร้างโฟลเดอร์ใหม่ใน Google Drive สำหรับเก็บรูป → กด "แชร์" →
     วาง client_email ลงไปเช่นกัน → ตั้งสิทธิ์เป็น "ผู้แก้ไข (Editor)"
   - คัดลอก Folder ID จาก URL (ส่วนหลัง /folders/) ใส่ในตัวแปร
     DRIVE_FOLDER_ID ด้านล่าง

3. ติดตั้ง library ที่ต้องใช้เพิ่ม (ใส่ใน requirements.txt):
   gspread
   google-auth
   google-api-python-client

โครงสร้างข้อมูลใน Google Sheet (สร้างให้อัตโนมัติ ไม่ต้องสร้างเอง):
  ชีต "Sites": SiteName | Status | CreatedAt | StartMile | EndMile | DataJSON
  (DataJSON เก็บข้อมูลโรงแรม/น้ำมันทั้งหมด รวม URL รูปจาก Drive)
═══════════════════════════════════════════════════════════════════════
"""

import gradio as gr
import pandas as pd
import json
import os
import io
from datetime import datetime
from docx import Document
from docx.shared import Inches
from PIL import Image as PILImage

import gspread
from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload

# ─────────────────────────────────────────────────────────────────────────
# 0. ตั้งค่า
# ─────────────────────────────────────────────────────────────────────────
KEY_FILE = os.environ.get("GOOGLE_KEY_FILE", "key.json")
def _get_secret(key, env_key=""):
    try:
        v = st.secrets.get(key, "")
        return v if v else ""
    except Exception:
        return os.environ.get(env_key or key, "")

SPREADSHEET_ID  = _get_secret("SPREADSHEET_ID",  "GOOGLE_SHEET_ID")
DRIVE_FOLDER_ID = _get_secret("DRIVE_FOLDER_ID", "GOOGLE_DRIVE_FOLDER_ID")

SCOPES = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive",
]
SHEET_NAME = "Sites"
SHEET_HEADER = ["SiteName", "Status", "CreatedAt", "StartMile", "EndMile", "DataJSON"]

HOTEL_COUNT = 3
HOTEL_ITEM_COUNT = 3
FUEL_COUNT = 20


# ─────────────────────────────────────────────────────────────────────────
# 1. Google API clients — สร้างครั้งเดียวแล้ว cache ไว้ใช้ซ้ำ
# ─────────────────────────────────────────────────────────────────────────
_creds = None
_gspread_client = None
_drive_service = None


def _get_credentials():
    """โหลด credentials:
    1. Streamlit Secrets [gcp_service_account]  (Streamlit Cloud)
    2. ไฟล์ key.json  (รันในเครื่อง)
    """
    global _creds
    if _creds is not None:
        return _creds
    try:
        sa_info = dict(st.secrets["gcp_service_account"])
        _creds = Credentials.from_service_account_info(sa_info, scopes=SCOPES)
        return _creds
    except (KeyError, FileNotFoundError):
        pass
    except Exception as e:
        raise RuntimeError(f"โหลด credentials จาก Streamlit Secrets ไม่สำเร็จ: {e}")
    if not os.path.exists(KEY_FILE):
        raise RuntimeError(
            "ไม่พบ credentials\n"
            "• Streamlit Cloud: ไปที่ App settings → Secrets ใส่ [gcp_service_account]\n"
            f"• Local: วางไฟล์ key.json ไว้ในโฟลเดอร์เดียวกับ app.py"
        )
    _creds = Credentials.from_service_account_file(KEY_FILE, scopes=SCOPES)
    return _creds


def _get_gspread_client():
    global _gspread_client
    if _gspread_client is None:
        _gspread_client = gspread.authorize(_get_credentials())
    return _gspread_client


def _get_drive_service():
    global _drive_service
    if _drive_service is None:
        _drive_service = build("drive", "v3", credentials=_get_credentials())
    return _drive_service


def _get_worksheet():
    """เปิด worksheet 'Sites' ในชีตที่ตั้งค่าไว้ — สร้างให้อัตโนมัติถ้ายังไม่มี"""
    if not SPREADSHEET_ID:
        raise RuntimeError(
            "ยังไม่ได้ตั้งค่า SPREADSHEET_ID — กรุณาใส่ Spreadsheet ID ของ "
            "Google Sheet ที่จะใช้เก็บข้อมูล (ดูวิธีหาใน comment หัวไฟล์)"
        )
    gc = _get_gspread_client()
    sh = gc.open_by_key(SPREADSHEET_ID)
    try:
        ws = sh.worksheet(SHEET_NAME)
    except gspread.WorksheetNotFound:
        ws = sh.add_worksheet(title=SHEET_NAME, rows=200, cols=len(SHEET_HEADER))
        ws.append_row(SHEET_HEADER)
    # เติม header ให้ถ้า worksheet มีอยู่แล้วแต่ว่างเปล่า (กันกรณีถูกล้างมาก่อน)
    if not ws.row_values(1):
        ws.update([SHEET_HEADER], value_input_option="RAW")
    return ws


# ─────────────────────────────────────────────────────────────────────────
# 2. การจัดการข้อมูลไซต์ — อ่าน/เขียน Google Sheet แทน CSV
# ─────────────────────────────────────────────────────────────────────────
def _empty_site_data() -> dict:
    """โครงสร้างข้อมูลเปล่าของไซต์ 1 ไซต์ (เก็บเป็น JSON ในคอลัมน์ DataJSON)"""
    return {
        "hotels": {
            str(h): {str(i): {"img": "", "desc": ""} for i in range(1, HOTEL_ITEM_COUNT + 1)}
            for h in range(1, HOTEL_COUNT + 1)
        },
        "car_img": "",
        "mile_start_img": "",
        "mile_end_img": "",
        "fuel": {
            str(n): {"bill": "", "pre": "", "post": "", "date": "", "province": ""}
            for n in range(1, FUEL_COUNT + 1)
        },
    }


def load_history() -> pd.DataFrame:
    """โหลดตารางสรุปไซต์ทั้งหมดจาก Google Sheet"""
    try:
        ws = _get_worksheet()
        records = ws.get_all_records()
        if not records:
            return pd.DataFrame(columns=SHEET_HEADER)
        df = pd.DataFrame(records)
        # เติมคอลัมน์ที่อาจขาด (เผื่อ schema เปลี่ยนทีหลัง) กัน KeyError
        for c in SHEET_HEADER:
            if c not in df.columns:
                df[c] = ""
        return df[SHEET_HEADER].astype(str)
    except Exception as e:
        print(f"[เตือน] โหลดข้อมูลจาก Google Sheet ไม่สำเร็จ: {e}")
        return pd.DataFrame(columns=SHEET_HEADER)


def _save_history(df: pd.DataFrame):
    """บันทึก DataFrame ทั้งก้อนกลับลง Google Sheet (เขียนทับทั้งชีต)"""
    ws = _get_worksheet()
    rows = [SHEET_HEADER] + df[SHEET_HEADER].astype(str).values.tolist()
    ws.clear()
    ws.update(rows, value_input_option="RAW")


def _load_site_data(site_name: str) -> dict:
    """โหลดข้อมูลเชิงลึกของไซต์ (โรงแรม/รถ/น้ำมัน) จากคอลัมน์ DataJSON"""
    df = load_history()
    row = df[df["SiteName"] == site_name]
    if row.empty:
        return _empty_site_data()
    raw = row.iloc[0].get("DataJSON", "")
    if not raw:
        return _empty_site_data()
    try:
        data = json.loads(raw)
        default = _empty_site_data()
        for k, v in default.items():
            data.setdefault(k, v)
        return data
    except Exception:
        return _empty_site_data()


def _save_site_data(site_name: str, site_data: dict):
    """บันทึกข้อมูลเชิงลึกของไซต์กลับลงคอลัมน์ DataJSON"""
    df = load_history()
    if site_name not in df["SiteName"].values:
        new_row = {
            "SiteName": site_name, "Status": "Draft",
            "CreatedAt": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M"),
            "StartMile": "0", "EndMile": "0",
            "DataJSON": json.dumps(site_data, ensure_ascii=False),
        }
        df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)
    else:
        idx = df[df["SiteName"] == site_name].index[0]
        df.at[idx, "DataJSON"] = json.dumps(site_data, ensure_ascii=False)
    _save_history(df)


def create_or_open_site(site_name: str):
    """สร้างไซต์ใหม่ถ้ายังไม่มี แล้วคืนชื่อไซต์ (ไว้ใช้เปิดหน้าบันทึกข้อมูล)"""
    site_name = (site_name or "").strip()
    if not site_name:
        raise gr.Error("กรุณาระบุชื่อไซต์งาน")

    df = load_history()
    if site_name not in df["SiteName"].values:
        new_row = {
            "SiteName": site_name, "Status": "Draft",
            "CreatedAt": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M"),
            "StartMile": "0", "EndMile": "0",
            "DataJSON": json.dumps(_empty_site_data(), ensure_ascii=False),
        }
        df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)
        _save_history(df)

    return site_name


# ─────────────────────────────────────────────────────────────────────────
# 3. อัปโหลดรูปขึ้น Google Drive ทันทีที่ผู้ใช้เลือก
# ─────────────────────────────────────────────────────────────────────────
_drive_folder_cache = {}  # {site_name: folder_id} กันค้นหาโฟลเดอร์ซ้ำทุกครั้ง


def _get_or_create_site_folder(site_name: str) -> str:
    """หาโฟลเดอร์ย่อยของไซต์นี้ใน Drive — สร้างให้ถ้ายังไม่มี คืนค่า folder ID"""
    if site_name in _drive_folder_cache:
        return _drive_folder_cache[site_name]
    if not DRIVE_FOLDER_ID:
        raise RuntimeError(
            "ยังไม่ได้ตั้งค่า DRIVE_FOLDER_ID — กรุณาใส่ Folder ID ของ "
            "Google Drive ที่จะใช้เก็บรูป (ดูวิธีหาใน comment หัวไฟล์)"
        )
    service = _get_drive_service()
    safe_name = "".join(c for c in site_name if c not in '\\/:*?"<>|').strip()

    # ค้นหาโฟลเดอร์ย่อยที่มีชื่อตรงกับไซต์นี้ก่อน (กันสร้างซ้ำ)
    query = (
        f"name = '{safe_name}' and mimeType = 'application/vnd.google-apps.folder' "
        f"and '{DRIVE_FOLDER_ID}' in parents and trashed = false"
    )
    results = service.files().list(q=query, fields="files(id, name)").execute()
    existing = results.get("files", [])
    if existing:
        folder_id = existing[0]["id"]
    else:
        meta = {
            "name": safe_name,
            "mimeType": "application/vnd.google-apps.folder",
            "parents": [DRIVE_FOLDER_ID],
        }
        created = service.files().create(body=meta, fields="id").execute()
        folder_id = created["id"]

    _drive_folder_cache[site_name] = folder_id
    return folder_id


def _upload_image(site_name: str, temp_filepath, dest_filename: str) -> str:
    """
    อัปโหลดรูปจาก temp path ของ Gradio ขึ้น Google Drive (โฟลเดอร์ย่อยของไซต์นี้)
    คืนค่า: URL สำหรับเปิดดูรูป (หรือ "" ถ้าไม่มีไฟล์/อัปโหลดไม่สำเร็จ)
    """
    if not temp_filepath:
        return ""
    try:
        service = _get_drive_service()
        folder_id = _get_or_create_site_folder(site_name)

        # แปลงรูปเป็น JPEG เสมอก่อนอัปโหลด (เผื่อไฟล์ต้นฉบับเป็น HEIC/WebP
        # ที่บางโปรแกรมเปิดไม่ได้ และเพื่อให้ไฟล์ขนาดเล็กลงสม่ำเสมอ)
        with PILImage.open(temp_filepath) as im:
            if im.mode in ("RGBA", "P", "LA"):
                im = im.convert("RGB")
            buf = io.BytesIO()
            im.save(buf, format="JPEG", quality=90)
            buf.seek(0)

        media = MediaIoBaseUpload(buf, mimetype="image/jpeg", resumable=False)
        meta = {"name": dest_filename + ".jpg", "parents": [folder_id]}

        # ถ้ามีไฟล์ชื่อนี้อยู่แล้วในโฟลเดอร์ (เคยอัปโหลดมาก่อน) ให้เขียนทับ
        # แทนการสร้างไฟล์ใหม่ซ้ำๆ ทุกครั้งที่แก้ไข
        query = (
            f"name = '{dest_filename}.jpg' and '{folder_id}' in parents "
            f"and trashed = false"
        )
        existing = service.files().list(q=query, fields="files(id)").execute().get("files", [])

        if existing:
            file_id = existing[0]["id"]
            service.files().update(fileId=file_id, media_body=media).execute()
        else:
            created = service.files().create(body=meta, media_body=media, fields="id").execute()
            file_id = created["id"]
            # เปิดสิทธิ์อ่านสาธารณะ (anyone with link) เพื่อให้ดูรูปได้ตรงๆ
            service.permissions().create(
                fileId=file_id, body={"type": "anyone", "role": "reader"}
            ).execute()

        return f"https://drive.google.com/uc?export=view&id={file_id}"
    except Exception as e:
        print(f"[เตือน] อัปโหลดรูปไม่สำเร็จ ({dest_filename}): {e}")
        return ""


def on_upload_hotel_img(site_name, hotel_no, item_no, filepath, site_data):
    """เรียกทันทีที่เลือก/เปลี่ยนรูปโรงแรม — อัปโหลด Drive ทันที ไม่ต้องกดบันทึก"""
    if not site_name:
        gr.Warning("กรุณาเลือกหรือสร้างไซต์งานก่อน")
        return site_data
    dest_url = _upload_image(site_name, filepath, f"hotel_{hotel_no}_{item_no}")
    if dest_url:
        site_data["hotels"][str(hotel_no)][str(item_no)]["img"] = dest_url
    return site_data


def on_upload_car_img(site_name, filepath, site_data):
    if not site_name:
        gr.Warning("กรุณาเลือกหรือสร้างไซต์งานก่อน")
        return site_data
    dest_url = _upload_image(site_name, filepath, "car")
    if dest_url:
        site_data["car_img"] = dest_url
    return site_data


def on_upload_mile_start(site_name, filepath, site_data):
    if not site_name:
        gr.Warning("กรุณาเลือกหรือสร้างไซต์งานก่อน")
        return site_data
    dest_url = _upload_image(site_name, filepath, "mile_start")
    if dest_url:
        site_data["mile_start_img"] = dest_url
    return site_data


def on_upload_mile_end(site_name, filepath, site_data):
    if not site_name:
        gr.Warning("กรุณาเลือกหรือสร้างไซต์งานก่อน")
        return site_data
    dest_url = _upload_image(site_name, filepath, "mile_end")
    if dest_url:
        site_data["mile_end_img"] = dest_url
    return site_data


def on_upload_fuel_img(site_name, fuel_no, kind, filepath, site_data):
    """kind: 'bill' | 'pre' | 'post'"""
    if not site_name:
        gr.Warning("กรุณาเลือกหรือสร้างไซต์งานก่อน")
        return site_data
    dest_url = _upload_image(site_name, filepath, f"fuel_{fuel_no}_{kind}")
    if dest_url:
        site_data["fuel"][str(fuel_no)][kind] = dest_url
    return site_data


# ─────────────────────────────────────────────────────────────────────────
# 4. ปุ่มบันทึกแยกแต่ละหัวข้อ
# ─────────────────────────────────────────────────────────────────────────
def save_hotel_section(site_name, hotel_no, site_data, desc1, desc2, desc3):
    if not site_name:
        raise gr.Error("กรุณาเลือกหรือสร้างไซต์งานก่อน")
    descs = [desc1, desc2, desc3]
    for i, d in enumerate(descs, start=1):
        site_data["hotels"][str(hotel_no)][str(i)]["desc"] = d or ""
    _save_site_data(site_name, site_data)
    return site_data, f"✅ บันทึกโรงแรมที่ {hotel_no} เรียบร้อย"


def save_car_section(site_name, mile_start, mile_end, site_data):
    if not site_name:
        raise gr.Error("กรุณาเลือกหรือสร้างไซต์งานก่อน")
    df = load_history()
    idx = df[df["SiteName"] == site_name].index
    if len(idx) > 0:
        df.at[idx[0], "StartMile"] = str(mile_start or 0)
        df.at[idx[0], "EndMile"] = str(mile_end or 0)
        _save_history(df)
    _save_site_data(site_name, site_data)
    return "✅ บันทึกข้อมูลรถเรียบร้อย"


def save_fuel_section(site_name, fuel_no, date_val, province, site_data):
    if not site_name:
        raise gr.Error("กรุณาเลือกหรือสร้างไซต์งานก่อน")
    site_data["fuel"][str(fuel_no)]["date"] = date_val or ""
    site_data["fuel"][str(fuel_no)]["province"] = province or ""
    _save_site_data(site_name, site_data)
    return site_data, f"✅ บันทึกการเติมน้ำมันครั้งที่ {fuel_no} เรียบร้อย"


# ─────────────────────────────────────────────────────────────────────────
# 5. ยืนยัน — สร้างไฟล์ Word ตามรูปแบบหน้าเว็บ แล้วเซฟลง Drive
# ─────────────────────────────────────────────────────────────────────────
def _fetch_drive_image_bytes(url: str):
    """ดาวน์โหลดรูปจาก Drive URL กลับมาเป็น bytes (สำหรับแทรกลง Word)"""
    if not url:
        return None
    try:
        import re
        m = re.search(r'[?&]id=([^&]+)', url)
        if not m:
            return None
        file_id = m.group(1)
        service = _get_drive_service()
        request = service.files().get_media(fileId=file_id)
        buf = io.BytesIO()
        from googleapiclient.http import MediaIoBaseDownload
        downloader = MediaIoBaseDownload(buf, request)
        done = False
        while not done:
            _, done = downloader.next_chunk()
        buf.seek(0)
        return buf.read()
    except Exception as e:
        print(f"[เตือน] ดาวน์โหลดรูปจาก Drive ไม่สำเร็จ ({url}): {e}")
        return None


def _safe_add_picture(target, img_url, width):
    """
    แทรกรูปจาก Drive URL ลง Word อย่างปลอดภัย — ดาวน์โหลดมาเป็น bytes ก่อน
    แล้วแปลงเป็น JPEG เสมอ (กัน python-docx อ่านฟอร์แมตแปลกๆ ไม่ออก)
    คืนค่า True ถ้าแทรกสำเร็จ, False ถ้าแทรกไม่ได้
    """
    if not img_url:
        return False
    raw = _fetch_drive_image_bytes(img_url)
    if not raw:
        return False
    try:
        with PILImage.open(io.BytesIO(raw)) as im:
            if im.mode in ("RGBA", "P", "LA"):
                im = im.convert("RGB")
            buf = io.BytesIO()
            im.save(buf, format="JPEG", quality=90)
            buf.seek(0)

        if hasattr(target, "add_picture"):
            target.add_picture(buf, width=width)
        else:
            p = target.add_paragraph()
            run = p.add_run()
            run.add_picture(buf, width=width)
        return True
    except Exception as e:
        print(f"[เตือน] แทรกรูปไม่สำเร็จ ({img_url}): {e}")
        return False


def confirm_and_export(site_name, site_data):
    if not site_name:
        raise gr.Error("กรุณาเลือกหรือสร้างไซต์งานก่อน")

    df = load_history()
    idx = df[df["SiteName"] == site_name].index
    mile_start = float(df.at[idx[0], "StartMile"]) if len(idx) > 0 and df.at[idx[0], "StartMile"] else 0
    mile_end = float(df.at[idx[0], "EndMile"]) if len(idx) > 0 and df.at[idx[0], "EndMile"] else 0

    doc = Document()
    doc.add_heading(f'Trip Report — {site_name}', 0)
    doc.add_paragraph(f"สร้างเมื่อ: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    # ── ส่วนที่ 1: ข้อมูลรถ ──
    doc.add_heading('ข้อมูลรถและการเดินทาง', level=1)
    distance = max(0, mile_end - mile_start)
    doc.add_paragraph(f"เลขไมล์เริ่มต้น: {mile_start:.0f}")
    doc.add_paragraph(f"เลขไมล์หลังจบ: {mile_end:.0f}")
    doc.add_paragraph(f"ระยะทางรวม: {distance:.0f} กม.")

    car_img = site_data.get("car_img", "")
    mile_start_img = site_data.get("mile_start_img", "")
    mile_end_img = site_data.get("mile_end_img", "")
    if car_img or mile_start_img or mile_end_img:
        car_table = doc.add_table(rows=1, cols=2)
        car_table.autofit = False
        car_table.columns[0].width = Inches(3.5)
        car_table.columns[1].width = Inches(3.0)

        left_cell = car_table.cell(0, 0)
        left_cell.paragraphs[0].text = "รูปรถ"
        if car_img:
            if not _safe_add_picture(left_cell, car_img, Inches(3.2)):
                left_cell.add_paragraph("(ไม่สามารถแสดงรูปนี้ได้)")
        else:
            left_cell.add_paragraph("(ไม่มีรูปรถ)")

        right_cell = car_table.cell(0, 1)
        right_cell.paragraphs[0].text = "ไมล์เริ่มต้น"
        if mile_start_img:
            if not _safe_add_picture(right_cell, mile_start_img, Inches(2.2)):
                right_cell.add_paragraph("(ไม่สามารถแสดงรูปนี้ได้)")
        else:
            right_cell.add_paragraph("(ไม่มีรูป)")
        right_cell.add_paragraph("ไมล์หลังจบ")
        if mile_end_img:
            if not _safe_add_picture(right_cell, mile_end_img, Inches(2.2)):
                right_cell.add_paragraph("(ไม่สามารถแสดงรูปนี้ได้)")
        else:
            right_cell.add_paragraph("(ไม่มีรูป)")
        doc.add_paragraph("")

    # ── ส่วนที่ 2: โรงแรม ──
    doc.add_heading('รูปภาพโรงแรม', level=1)
    any_hotel = False
    for h in range(1, HOTEL_COUNT + 1):
        items = site_data["hotels"][str(h)]
        has_any_img = any(items[str(i)]["img"] for i in range(1, HOTEL_ITEM_COUNT + 1))
        if not has_any_img:
            continue
        any_hotel = True
        doc.add_heading(f'โรงแรมที่ {h}', level=2)

        hotel_table = doc.add_table(rows=2, cols=HOTEL_ITEM_COUNT)
        hotel_table.autofit = True
        img_row = hotel_table.rows[0]
        desc_row = hotel_table.rows[1]

        for i in range(1, HOTEL_ITEM_COUNT + 1):
            item = items[str(i)]
            img_cell = img_row.cells[i - 1]
            desc_cell = desc_row.cells[i - 1]
            if item["img"]:
                if not _safe_add_picture(img_cell, item["img"], Inches(1.9)):
                    img_cell.paragraphs[0].text = "(ไม่สามารถแสดงรูปนี้ได้)"
            else:
                img_cell.paragraphs[0].text = "(ไม่มีรูป)"
            desc_cell.paragraphs[0].text = item["desc"] or "(ไม่มีคำอธิบาย)"
        doc.add_paragraph("")
    if not any_hotel:
        doc.add_paragraph("(ยังไม่มีรูปโรงแรมที่บันทึกแล้ว)")

    # ── ส่วนที่ 3: น้ำมัน ──
    doc.add_heading('บันทึกการเติมน้ำมัน', level=1)
    any_fuel = False
    for n in range(1, FUEL_COUNT + 1):
        f = site_data["fuel"][str(n)]
        if not (f["bill"] or f["pre"] or f["post"]):
            continue
        any_fuel = True
        doc.add_heading(f'การเติมครั้งที่ {n}', level=2)
        doc.add_paragraph(f"วันที่: {f['date'] or '-'}   จังหวัด: {f['province'] or '-'}")

        fuel_table = doc.add_table(rows=1, cols=2)
        fuel_table.autofit = False
        fuel_table.columns[0].width = Inches(4.0)
        fuel_table.columns[1].width = Inches(2.5)

        left_cell = fuel_table.cell(0, 0)
        left_cell.paragraphs[0].text = "ใบเสร็จ"
        if f["bill"]:
            if not _safe_add_picture(left_cell, f["bill"], Inches(3.7)):
                left_cell.add_paragraph("(ไม่สามารถแสดงรูปนี้ได้)")
        else:
            left_cell.add_paragraph("(ไม่มีรูปใบเสร็จ)")

        right_cell = fuel_table.cell(0, 1)
        right_cell.paragraphs[0].text = "ไมล์ก่อนเติม"
        if f["pre"]:
            if not _safe_add_picture(right_cell, f["pre"], Inches(2.2)):
                right_cell.add_paragraph("(ไม่สามารถแสดงรูปนี้ได้)")
        else:
            right_cell.add_paragraph("(ไม่มีรูป)")
        right_cell.add_paragraph("ไมล์หลังเติม")
        if f["post"]:
            if not _safe_add_picture(right_cell, f["post"], Inches(2.2)):
                right_cell.add_paragraph("(ไม่สามารถแสดงรูปนี้ได้)")
        else:
            right_cell.add_paragraph("(ไม่มีรูป)")

        doc.add_paragraph("")
    if not any_fuel:
        doc.add_paragraph("(ยังไม่มีรายการเติมน้ำมันที่บันทึกแล้ว)")

    # บันทึกไฟล์ Word ลง Google Drive (โฟลเดอร์ของไซต์นี้)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    report_url = ""
    try:
        service = _get_drive_service()
        folder_id = _get_or_create_site_folder(site_name)
        media = MediaIoBaseUpload(
            buf,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            resumable=False,
        )
        meta = {"name": f"{site_name}_Report.docx", "parents": [folder_id]}

        query = (
            f"name = '{site_name}_Report.docx' and '{folder_id}' in parents "
            f"and trashed = false"
        )
        existing = service.files().list(q=query, fields="files(id)").execute().get("files", [])
        if existing:
            file_id = existing[0]["id"]
            service.files().update(fileId=file_id, media_body=media).execute()
        else:
            created = service.files().create(body=meta, media_body=media, fields="id").execute()
            file_id = created["id"]
            service.permissions().create(
                fileId=file_id, body={"type": "anyone", "role": "reader"}
            ).execute()
        report_url = f"https://drive.google.com/file/d/{file_id}/view"
    except Exception as e:
        print(f"[เตือน] อัปโหลดไฟล์รายงานไม่สำเร็จ: {e}")

    # อัปเดตสถานะเป็น Confirmed
    df = load_history()
    idx2 = df[df["SiteName"] == site_name].index
    if len(idx2) > 0:
        df.at[idx2[0], "Status"] = "Confirmed"
        _save_history(df)

    if report_url:
        return f"✅ ยืนยันสำเร็จ — ไฟล์รายงานอยู่ที่:\n{report_url}", load_history()
    return "⚠️ ยืนยันสำเร็จ แต่อัปโหลดไฟล์รายงานไม่สำเร็จ (ดู log สำหรับรายละเอียด)", load_history()


# ─────────────────────────────────────────────────────────────────────────
# 6. โหลดข้อมูลไซต์เข้า UI ตอนกด "แก้ไข" จากตารางประวัติ
# ─────────────────────────────────────────────────────────────────────────
def get_selected_row(evt: gr.SelectData):
    df = load_history()
    row = df.iloc[evt.index[0]]
    return row["SiteName"]


def open_site_for_edit(site_name):
    """โหลดข้อมูลทั้งหมดของไซต์ที่เลือก คืนค่าให้ทุก component ในหน้าบันทึกข้อมูล"""
    site_name = (site_name or "").strip()
    if not site_name:
        raise gr.Error("กรุณาเลือกไซต์งานจากตารางก่อน")

    site_name = create_or_open_site(site_name)
    site_data = _load_site_data(site_name)

    df = load_history()
    row = df[df["SiteName"] == site_name].iloc[0]
    mile_start_raw = row.get("StartMile", "")
    mile_end_raw = row.get("EndMile", "")
    mile_start = float(mile_start_raw) if mile_start_raw else 0
    mile_end = float(mile_end_raw) if mile_end_raw else 0

    outputs = [site_name, site_data, gr.Tabs(selected="entry")]

    # โรงแรม: 9 รูป (เป็น URL — gr.Image แสดงผลจาก URL ได้โดยตรง) + 9 คำอธิบาย
    for h in range(1, HOTEL_COUNT + 1):
        for i in range(1, HOTEL_ITEM_COUNT + 1):
            item = site_data["hotels"][str(h)][str(i)]
            img_val = item["img"] if item["img"] else None
            outputs.append(img_val)
            outputs.append(item["desc"])

    car_img = site_data["car_img"] or None
    mile_start_img = site_data["mile_start_img"] or None
    mile_end_img = site_data["mile_end_img"] or None
    outputs += [car_img, mile_start_img, mile_end_img, mile_start, mile_end]

    for n in range(1, FUEL_COUNT + 1):
        f = site_data["fuel"][str(n)]
        outputs += [f["bill"] or None, f["pre"] or None, f["post"] or None, f["date"], f["province"]]

    return outputs


# ─────────────────────────────────────────────────────────────────────────
# 7. สร้าง UI
# ─────────────────────────────────────────────────────────────────────────
with gr.Blocks(title="ระบบจัดการไซต์งาน") as demo:
    site_data_state = gr.State(_empty_site_data())

    with gr.Tabs() as tabs:
        with gr.Tab("หน้าแรก: รายการไซต์งาน", id="home"):
            gr.Markdown("### ประวัติไซต์งาน (คลิกที่แถวในตารางเพื่อเลือก)")
            _initial_history = load_history()
            site_list = gr.Dataframe(
                value=_initial_history[["SiteName", "Status", "CreatedAt"]],
                headers=["ชื่อไซต์", "สถานะ", "วันที่สร้าง"],
                interactive=False,
            )
            selected_site_name = gr.Textbox(visible=False)
            edit_btn = gr.Button("✏️ แก้ไขไซต์ที่เลือก", variant="secondary")

            gr.Markdown("### หรือสร้างไซต์งานใหม่")
            with gr.Row():
                new_site_input = gr.Textbox(label="ชื่อไซต์งานใหม่", scale=3)
                create_btn = gr.Button("➕ สร้าง", variant="primary", scale=1)

        with gr.Tab("บันทึกข้อมูล", id="entry"):
            back_btn = gr.Button("⬅️ กลับหน้าแรก")
            site_name_display = gr.Textbox(label="ไซต์งานปัจจุบัน", interactive=False)

            hotel_imgs = {}
            hotel_descs = {}
            hotel_save_btns = {}
            hotel_status = {}

            with gr.Accordion("1. ข้อมูลโรงแรม", open=False):
                for h in range(1, HOTEL_COUNT + 1):
                    gr.Markdown(f"#### โรงแรมที่ {h}")
                    with gr.Row():
                        for i in range(1, HOTEL_ITEM_COUNT + 1):
                            with gr.Column():
                                img = gr.Image(label=f"รูป {i}", type="filepath")
                                desc = gr.Textbox(label=f"คำอธิบายรูป {i}", lines=2)
                                hotel_imgs[(h, i)] = img
                                hotel_descs[(h, i)] = desc
                    hotel_status[h] = gr.Textbox(label="", interactive=False, show_label=False)
                    hotel_save_btns[h] = gr.Button(f"💾 บันทึกโรงแรมที่ {h}")

            with gr.Accordion("2. ข้อมูลรถ", open=False):
                with gr.Row():
                    car_img_input = gr.Image(label="รูปรถ", type="filepath", scale=2)
                    with gr.Column(scale=1):
                        mile_start_img_input = gr.Image(label="รูปไมล์เริ่ม", type="filepath")
                        mile_end_img_input = gr.Image(label="รูปไมล์จบ", type="filepath")
                with gr.Row():
                    mile_start_input = gr.Number(label="เลขไมล์เริ่มต้น", value=0)
                    mile_end_input = gr.Number(label="เลขไมล์หลังจบ", value=0)
                car_status = gr.Textbox(label="", interactive=False, show_label=False)
                save_car_btn = gr.Button("💾 บันทึกข้อมูลรถ")

            fuel_bill = {}
            fuel_pre = {}
            fuel_post = {}
            fuel_date = {}
            fuel_province = {}
            fuel_save_btns = {}
            fuel_status = {}

            with gr.Accordion("3. การเติมน้ำมัน (20 ครั้ง)", open=False):
                for n in range(1, FUEL_COUNT + 1):
                    with gr.Accordion(f"เติมครั้งที่ {n}", open=False):
                        with gr.Row():
                            bill = gr.Image(label="ใบเสร็จ (รูปใหญ่)", type="filepath", scale=2)
                            with gr.Column(scale=1):
                                pre = gr.Image(label="ไมล์ก่อนเติม", type="filepath")
                                post = gr.Image(label="ไมล์หลังเติม", type="filepath")
                        with gr.Row():
                            date_box = gr.Textbox(label="วันที่เติม (เช่น 2026-06-21)")
                            prov_box = gr.Textbox(label="จังหวัด")
                        fuel_status[n] = gr.Textbox(label="", interactive=False, show_label=False)
                        fuel_save_btns[n] = gr.Button(f"💾 บันทึกการเติมครั้งที่ {n}")

                        fuel_bill[n] = bill
                        fuel_pre[n] = pre
                        fuel_post[n] = post
                        fuel_date[n] = date_box
                        fuel_province[n] = prov_box

            gr.Markdown("---")
            export_btn = gr.Button("✅ ยืนยันและสร้างไฟล์ Word (เซฟลง Drive)", variant="primary")
            export_status = gr.Textbox(label="สถานะ")

    # ═══════════════════════════════════════════════════════════════════
    # Wiring เหตุการณ์ทั้งหมด
    # ═══════════════════════════════════════════════════════════════════
    site_list.select(get_selected_row, None, selected_site_name)

    all_entry_outputs = [site_name_display, site_data_state, tabs]
    for h in range(1, HOTEL_COUNT + 1):
        for i in range(1, HOTEL_ITEM_COUNT + 1):
            all_entry_outputs += [hotel_imgs[(h, i)], hotel_descs[(h, i)]]
    all_entry_outputs += [car_img_input, mile_start_img_input, mile_end_img_input,
                           mile_start_input, mile_end_input]
    for n in range(1, FUEL_COUNT + 1):
        all_entry_outputs += [fuel_bill[n], fuel_pre[n], fuel_post[n], fuel_date[n], fuel_province[n]]

    edit_btn.click(open_site_for_edit, inputs=selected_site_name, outputs=all_entry_outputs)

    create_btn.click(
        lambda name: create_or_open_site(name),
        inputs=new_site_input, outputs=selected_site_name
    ).then(
        open_site_for_edit, inputs=selected_site_name, outputs=all_entry_outputs
    )

    back_btn.click(
        lambda: (gr.Tabs(selected="home"), load_history()[["SiteName", "Status", "CreatedAt"]]),
        None, [tabs, site_list]
    )

    for (h, i), img_comp in hotel_imgs.items():
        img_comp.change(
            lambda filepath, name, sd, h=h, i=i: on_upload_hotel_img(name, h, i, filepath, sd),
            inputs=[img_comp, site_name_display, site_data_state], outputs=site_data_state
        )

    for h in range(1, HOTEL_COUNT + 1):
        d1, d2, d3 = hotel_descs[(h, 1)], hotel_descs[(h, 2)], hotel_descs[(h, 3)]
        hotel_save_btns[h].click(
            save_hotel_section,
            inputs=[site_name_display, gr.State(h), site_data_state, d1, d2, d3],
            outputs=[site_data_state, hotel_status[h]]
        )

    car_img_input.change(on_upload_car_img, inputs=[site_name_display, car_img_input, site_data_state], outputs=site_data_state)
    mile_start_img_input.change(on_upload_mile_start, inputs=[site_name_display, mile_start_img_input, site_data_state], outputs=site_data_state)
    mile_end_img_input.change(on_upload_mile_end, inputs=[site_name_display, mile_end_img_input, site_data_state], outputs=site_data_state)

    save_car_btn.click(
        save_car_section,
        inputs=[site_name_display, mile_start_input, mile_end_input, site_data_state],
        outputs=car_status
    )

    for n in range(1, FUEL_COUNT + 1):
        fuel_bill[n].change(
            lambda filepath, name, sd, n=n: on_upload_fuel_img(name, n, "bill", filepath, sd),
            inputs=[fuel_bill[n], site_name_display, site_data_state], outputs=site_data_state
        )
        fuel_pre[n].change(
            lambda filepath, name, sd, n=n: on_upload_fuel_img(name, n, "pre", filepath, sd),
            inputs=[fuel_pre[n], site_name_display, site_data_state], outputs=site_data_state
        )
        fuel_post[n].change(
            lambda filepath, name, sd, n=n: on_upload_fuel_img(name, n, "post", filepath, sd),
            inputs=[fuel_post[n], site_name_display, site_data_state], outputs=site_data_state
        )
        fuel_save_btns[n].click(
            save_fuel_section,
            inputs=[site_name_display, gr.State(n), fuel_date[n], fuel_province[n], site_data_state],
            outputs=[site_data_state, fuel_status[n]]
        )

    export_btn.click(
        confirm_and_export,
        inputs=[site_name_display, site_data_state],
        outputs=[export_status, site_list]
    )

if __name__ == "__main__":
    demo.launch(server_name="0.0.0.0", server_port=int(os.environ.get("PORT", 7860)))
